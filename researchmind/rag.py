from __future__ import annotations

import json
import math
from pathlib import Path
import re
import shutil
from uuid import uuid4
from typing import Iterable

from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from researchmind import config


class RAGService:
    def __init__(self) -> None:
        config.ensure_dirs()
        self._parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.PARENT_CHUNK_SIZE,
            chunk_overlap=config.PARENT_CHUNK_OVERLAP,
        )
        self._child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHILD_CHUNK_SIZE,
            chunk_overlap=config.CHILD_CHUNK_OVERLAP,
        )
        self._embeddings: HuggingFaceEmbeddings | None = None
        self._vector_store: Chroma | None = None

    def _get_embeddings(self) -> HuggingFaceEmbeddings:
        if self._embeddings is None:
            self._embeddings = HuggingFaceEmbeddings(model_name=config.EMBEDDING_MODEL)
        return self._embeddings

    def _get_vector_store(self) -> Chroma:
        if self._vector_store is None:
            try:
                self._vector_store = Chroma(
                    collection_name="researchmind_docs",
                    embedding_function=self._get_embeddings(),
                    persist_directory=str(config.CHROMA_DIR),
                )
            except Exception as error:
                error_text = str(error).lower()
                recoverable_markers = (
                    "default_tenant",
                    "tenant",
                    "rustbindingsapi",
                    "has no attribute 'bindings'",
                    'has no attribute "bindings"',
                    "keyerror: '_type'",
                    "'_type'",
                )
                if not any(marker in error_text for marker in recoverable_markers):
                    raise

                if config.CHROMA_DIR.exists():
                    shutil.rmtree(config.CHROMA_DIR, ignore_errors=True)
                config.CHROMA_DIR.mkdir(parents=True, exist_ok=True)

                try:
                    from chromadb.api.client import SharedSystemClient

                    SharedSystemClient.clear_system_cache()
                except Exception:
                    pass

                try:
                    self._vector_store = Chroma(
                        collection_name="researchmind_docs",
                        embedding_function=self._get_embeddings(),
                        persist_directory=str(config.CHROMA_DIR),
                    )
                except Exception as retry_error:
                    retry_text = str(retry_error).lower()
                    if "rustbindingsapi" in retry_text or "has no attribute 'bindings'" in retry_text:
                        raise RuntimeError(
                            "检测到 ChromaDB 版本不兼容（RustBindingsAPI.bindings）。"
                            "请将 chromadb 降级到 <1.0（例如 chromadb>=0.5.23,<1.0.0），"
                            "然后重装依赖并重启应用。"
                        ) from retry_error
                    raise
        return self._vector_store

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        return filename.replace("/", "_").replace("\\", "_")

    def _save_local_file(self, filename: str, content: bytes) -> Path:
        safe_name = self._sanitize_filename(filename)
        file_path = config.UPLOAD_DIR / safe_name
        file_path.write_bytes(content)
        return file_path

    def _load_parent_store(self) -> dict[str, dict]:
        if not config.PARENT_STORE_FILE.exists():
            return {}
        try:
            payload = json.loads(config.PARENT_STORE_FILE.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
        if not isinstance(payload, dict):
            return {}
        return payload

    def _save_parent_store(self, parent_store: dict[str, dict]) -> None:
        config.PARENT_STORE_FILE.write_text(
            json.dumps(parent_store, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def list_uploaded_files(self) -> list[str]:
        if not config.UPLOAD_DIR.exists():
            return []
        files = sorted(
            [
                path.name
                for path in config.UPLOAD_DIR.iterdir()
                if path.is_file() and path.suffix.lower() in {".pdf", ".docx"}
            ]
        )
        return files

    def _load_web_sources(self) -> list[str]:
        if not config.WEB_SOURCES_FILE.exists():
            return []
        try:
            data = json.loads(config.WEB_SOURCES_FILE.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return []
        if not isinstance(data, list):
            return []
        return sorted([str(item) for item in data if isinstance(item, str)])

    def _save_web_sources(self, urls: list[str]) -> None:
        config.WEB_SOURCES_FILE.write_text(
            json.dumps(sorted(set(urls)), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def list_knowledge_sources(self) -> list[str]:
        local_files = self.list_uploaded_files()
        web_sources = [f"[网页] {url}" for url in self._load_web_sources()]
        return local_files + web_sources

    def _load_pdf(self, file_path: Path) -> list[Document]:
        reader = PdfReader(str(file_path))
        pages: list[Document] = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            pages.append(
                Document(
                    page_content=text,
                    metadata={"source": file_path.name, "page": page_index},
                )
            )
        return pages

    def _load_docx(self, file_path: Path) -> list[Document]:
        doc = DocxDocument(str(file_path))
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        if not paragraphs:
            return []
        return [
            Document(
                page_content="\n".join(paragraphs),
                metadata={"source": file_path.name, "page": 1},
            )
        ]

    def _load_web_url(self, url: str) -> list[Document]:
        response = requests.get(url, timeout=20)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        title = soup.title.string.strip() if soup.title and soup.title.string else "Untitled"
        content = "\n".join(line.strip() for line in soup.get_text(separator="\n").splitlines() if line.strip())
        if not content:
            return []

        return [
            Document(
                page_content=content,
                metadata={"source": url, "page": 1, "title": title},
            )
        ]

    def _ingest_documents(self, docs: list[Document]) -> int:
        child_chunks: list[Document] = []
        parent_store = self._load_parent_store()

        for document in docs:
            source = str(document.metadata.get("source", "unknown.source"))
            page = int(document.metadata.get("page", 1))
            text = (document.page_content or "").strip()
            if not text:
                continue

            parent_chunks = self._parent_splitter.split_text(text)
            if not parent_chunks:
                continue

            for parent_index, parent_text in enumerate(parent_chunks, start=1):
                clean_parent = parent_text.strip()
                if not clean_parent:
                    continue
                parent_id = f"{source}::p{page}::pi{parent_index}::{uuid4().hex[:10]}"
                parent_store[parent_id] = {
                    "text": clean_parent,
                    "metadata": {
                        "source": source,
                        "page": page,
                        "parent_id": parent_id,
                        "parent_index": parent_index,
                    },
                }

                sub_chunks = self._child_splitter.split_text(clean_parent)
                for child_index, child_text in enumerate(sub_chunks, start=1):
                    clean_child = child_text.strip()
                    if not clean_child:
                        continue
                    chunk_id = f"{parent_id}::c{child_index}"
                    child_chunks.append(
                        Document(
                            page_content=clean_child,
                            metadata={
                                "source": source,
                                "page": page,
                                "parent_id": parent_id,
                                "parent_index": parent_index,
                                "child_index": child_index,
                                "chunk_id": chunk_id,
                            },
                        )
                    )

        if not child_chunks:
            return 0

        self._save_parent_store(parent_store)
        self._get_vector_store().add_documents(child_chunks)
        return len(child_chunks)

    def ingest_file(self, filename: str, content: bytes) -> int:
        file_path = self._save_local_file(filename, content)
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            docs = self._load_pdf(file_path)
        elif suffix == ".docx":
            docs = self._load_docx(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {suffix}，当前仅支持 .pdf 和 .docx")

        return self._ingest_documents(docs)

    def ingest_web_url(self, url: str) -> int:
        normalized_url = url.strip()
        if not normalized_url:
            raise ValueError("网页 URL 不能为空")
        if not normalized_url.startswith(("http://", "https://")):
            normalized_url = f"https://{normalized_url}"

        docs = self._load_web_url(normalized_url)
        chunk_count = self._ingest_documents(docs)

        existing = self._load_web_sources()
        if normalized_url not in existing:
            existing.append(normalized_url)
            self._save_web_sources(existing)

        return chunk_count

    def has_indexed_docs(self) -> bool:
        return len(self.list_knowledge_sources()) > 0

    @staticmethod
    def _tokenize(text: str) -> set[str]:
        return {token for token in re.findall(r"[\w\u4e00-\u9fa5]+", (text or "").lower()) if token}

    @classmethod
    def _sparse_score(cls, query: str, text: str) -> float:
        query_tokens = cls._tokenize(query)
        text_tokens = cls._tokenize(text)
        if not query_tokens or not text_tokens:
            return 0.0
        overlap = len(query_tokens & text_tokens)
        if overlap == 0:
            return 0.0
        return overlap / math.sqrt(len(text_tokens) + 1)

    def _dense_candidates(self, query: str, candidate_k: int) -> list[tuple[Document, float]]:
        vector_store = self._get_vector_store()
        try:
            pairs = vector_store.similarity_search_with_relevance_scores(query, k=candidate_k)
            return [(doc, float(score)) for doc, score in pairs]
        except Exception:
            docs = vector_store.similarity_search(query, k=candidate_k)
            return [(doc, 0.5) for doc in docs]

    def _sparse_parent_candidates(self, query: str, candidate_k: int) -> list[tuple[str, float]]:
        parent_store = self._load_parent_store()
        scored: list[tuple[str, float]] = []
        for parent_id, payload in parent_store.items():
            text = str(payload.get("text", ""))
            score = self._sparse_score(query, text)
            if score > 0:
                scored.append((parent_id, score))
        scored.sort(key=lambda item: item[1], reverse=True)
        return scored[:candidate_k]

    @staticmethod
    def _normalize_scores(pairs: list[tuple[str, float]]) -> dict[str, float]:
        if not pairs:
            return {}
        values = [score for _, score in pairs]
        max_score = max(values)
        min_score = min(values)
        if math.isclose(max_score, min_score):
            return {key: 1.0 for key, _ in pairs}
        return {key: (score - min_score) / (max_score - min_score) for key, score in pairs}

    def retrieve(self, query: str, top_k: int = config.TOP_K) -> list[Document]:
        if not query.strip():
            return []

        candidate_k = max(top_k, config.HYBRID_CANDIDATE_K)
        dense_hits = self._dense_candidates(query, candidate_k)

        dense_by_parent: dict[str, float] = {}
        for doc, score in dense_hits:
            parent_id = str(doc.metadata.get("parent_id", "")).strip()
            if not parent_id:
                continue
            dense_by_parent[parent_id] = max(dense_by_parent.get(parent_id, 0.0), float(score))

        sparse_hits = self._sparse_parent_candidates(query, candidate_k)
        sparse_by_parent = dict(sparse_hits)

        normalized_dense = self._normalize_scores(list(dense_by_parent.items()))
        normalized_sparse = self._normalize_scores(list(sparse_by_parent.items()))

        all_parent_ids = set(normalized_dense.keys()) | set(normalized_sparse.keys())
        if not all_parent_ids:
            return []

        parent_store = self._load_parent_store()
        hybrid_scores: list[tuple[str, float]] = []
        for parent_id in all_parent_ids:
            dense_score = normalized_dense.get(parent_id, 0.0)
            sparse_score = normalized_sparse.get(parent_id, 0.0)
            hybrid_score = config.HYBRID_WEIGHT_DENSE * dense_score + config.HYBRID_WEIGHT_SPARSE * sparse_score
            hybrid_scores.append((parent_id, hybrid_score))

        reranked: list[tuple[str, float]] = []
        for parent_id, hybrid_score in hybrid_scores:
            payload = parent_store.get(parent_id, {})
            parent_text = str(payload.get("text", ""))
            lexical = self._sparse_score(query, parent_text)
            rerank_score = 0.7 * hybrid_score + 0.3 * lexical
            reranked.append((parent_id, rerank_score))

        reranked.sort(key=lambda item: item[1], reverse=True)
        limited = reranked[: max(top_k, config.RERANK_MAX_DOCS)]

        documents: list[Document] = []
        for parent_id, final_score in limited:
            payload = parent_store.get(parent_id)
            if not payload:
                continue
            metadata = dict(payload.get("metadata", {}))
            metadata.update({
                "retrieval_score": round(final_score, 6),
                "retrieval_mode": "hybrid_rerank",
            })
            documents.append(Document(page_content=str(payload.get("text", "")), metadata=metadata))
            if len(documents) >= top_k:
                break

        return documents

    @staticmethod
    def format_context(documents: Iterable[Document]) -> str:
        blocks: list[str] = []
        for idx, document in enumerate(documents, start=1):
            source = document.metadata.get("source", "unknown.source")
            page = document.metadata.get("page", "?")
            score = document.metadata.get("retrieval_score", "-")
            mode = document.metadata.get("retrieval_mode", "dense")
            blocks.append(f"[{idx}] 来源: {source} | 页码: {page} | 分数: {score} | 模式: {mode}\n{document.page_content}")
        return "\n\n".join(blocks)
